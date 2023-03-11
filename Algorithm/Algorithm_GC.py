from GeneralCorrosionMethod import GeneralCorrosionMethod as g
import os
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

# Set the directory where the file should be saved
directory = r"..\reports"

# Create the document
document = Document()

# Set the font to one suitable for mathematical expressions
math_font = document.styles.add_style('MathFont', WD_STYLE_TYPE.PARAGRAPH)
math_font.font.name = 'Cambria Math'
math_font.font.size = Pt(12)



# input data
t = 3.7592
pipe_type = 'Seamless'
M_ut = 0
FCA_ml = 0.14
LOSS = 0
FCA = 0
type_of_wall_loss = 'external'
t_mm = 3.13
RSF_a = 0.9

P = 30
S = 137.9
D_0 = 150
E = 1
Y_B31 = 0.4
MA = 0
t_sl = 0
t_amS = 3.13
t_amC = 3.13

method = 'general corrosion'

input_data_text = """
Input data:
t = 3.7592
pipe_type = 'Seamless'
M_ut = 0
FCA_ml = 0.14
LOSS = 0
FCA = 0
type_of_wall_loss = 'external'
t_mm = 3.13
RSF_a = 0.9

P = 30
S = 137.9
D_0 = 150
E = 1
Y_B31 = 0.4
MA = 0
t_sl = 0
t_amS = 3.13
t_amC = 3.13
"""

p = document.add_paragraph(input_data_text)
p.style = math_font

# algorithm
if method == 'general corrosion':
    # step 1 ----------------------------------------------------------------
    p = document.add_paragraph('STEP 1')

    # Nominal or furnished thickness of the component adjusted for mill undertolerance as applicable, [mm]
    t_nom = g.f_t_nom(pipe_type, t, M_ut).result
    p = document.add_paragraph(g.f_t_nom(pipe_type, t, M_ut).text)
    p.style = math_font

    # Nominal thickness in the region of corrosion corrected FCA_ml, [mm]
    t_ml = g.f_t_ml(t_nom, FCA_ml).result
    p = document.add_paragraph(g.f_t_ml(t_nom, FCA_ml).text)
    p.style = math_font

    # step 2 ----------------------------------------------------------------
    p = document.add_paragraph('STEP 2')

    # Future corroded wall thickness away from the damage area, [mm]
    t_c = g.f_t_c(t_nom, LOSS, FCA).result
    p = document.add_paragraph(g.f_t_c(t_nom, LOSS, FCA).text)
    p.style = math_font

    # step 3 ----------------------------------------------------------------
    p = document.add_paragraph('STEP 3')

    # Inside diameter of the cylinder, cone (at the location of the flaw), [mm]
    D = g.f_D(D_0, t_nom).result
    p = document.add_paragraph(g.f_D(D_0, t_nom).text)
    p.style = math_font

    # Inside diameter of the cylinder corrected for FCA_ml, [mm]
    D_ml = g.f_D_ml(type_of_wall_loss, D, FCA_ml).result
    p = document.add_paragraph(g.f_D_ml(type_of_wall_loss, D, FCA_ml).text)
    p.style = math_font

    # step 4 ----------------------------------------------------------------
    p = document.add_paragraph('STEP 4')

    # Remaining thickness ratio R_t, [-]
    R_t = g.f_R_t(t_mm, FCA_ml, t_ml).result
    p = document.add_paragraph(g.f_R_t(t_mm, FCA_ml, t_ml).text)
    p.style = math_font

    # step 5 ----------------------------------------------------------------
    p = document.add_paragraph('STEP 5')

    # Factor used to determine the length for thickness averaging based on an allowable Remaining Strength Factor and the remaining thickness ratio Rt, [-]
    Q = g.f_Q(R_t, RSF_a).result
    p = document.add_paragraph(g.f_Q(R_t, RSF_a).text)
    p.style = math_font

    # step 6 ----------------------------------------------------------------
    p = document.add_paragraph('STEP 6')

    #  Length for thickness averaging L, [mm]
    L = g.f_L(Q, D_ml, t_ml).result
    p = document.add_paragraph(g.f_L(Q, D_ml, t_ml).text)
    p.style = math_font

    # step 7 ----------------------------------------------------------------
    p = document.add_paragraph('STEP 7')

    # Minimum required thickness based on the circumferential stress, [mm]
    t_minC = g.f_t_minC(P, D_0, S, E, Y_B31, MA).result
    p = document.add_paragraph(g.f_t_minC(P, D_0, S, E, Y_B31, MA).text)
    p.style = math_font

    #  Minimum required thickness based on the longitudinal stress, [mm]
    t_minL = g.f_t_minL(P, D_0, S, E, Y_B31, MA, t_sl).result
    p = document.add_paragraph(g.f_t_minL(P, D_0, S, E, Y_B31, MA, t_sl).text)
    p.style = math_font

    # Minimum required thickness , [mm]
    t_min = g.f_t_min(t_minC, t_minL).result
    p = document.add_paragraph( g.f_t_min(t_minC, t_minL).text)
    p.style = math_font

    # step 8 ----------------------------------------------------------------
    p = document.add_paragraph('STEP 8')

    # Maximum allowable working pressure based on circumferential stress, [bar]
    MAWP_C = g.f_MAWP_C(S, E, t_c, MA, D_0, Y_B31).result
    p = document.add_paragraph(g.f_MAWP_C(S, E, t_c, MA, D_0, Y_B31).text)
    p.style = math_font

    #  Maximum allowable working pressure based on longitudinal stress, [bar]
    MAWP_L = g.f_MAWP_L(S, E, t_c, MA, D_0, Y_B31, t_sl).result
    p = document.add_paragraph(g.f_MAWP_L(S, E, t_c, MA, D_0, Y_B31, t_sl).text)
    p.style = math_font

    # Maximum allowable working pressure, [bar]
    MAWP = g.f_MAWP(MAWP_C, MAWP_L).result
    p = document.add_paragraph(g.f_MAWP(MAWP_C, MAWP_L).text)
    p.style = math_font

    # step 9 ----------------------------------------------------------------
    p = document.add_paragraph('STEP 9')

    # Average Measured Thickness from Critical Thickness Profiles based on the
    # longitudinal CTP determined at the time of the inspection
    average_longitudinal_thickness_criteria = g.check_average_longitudinal_thickness_criteria(t_amS, FCA_ml,
                                                                                              t_minC).result
    p = document.add_paragraph(g.check_average_longitudinal_thickness_criteria(t_amS, FCA_ml,
                                                                                              t_minC).text)
    p.style = math_font

    if average_longitudinal_thickness_criteria == 'failed':

        print(average_longitudinal_thickness_criteria)

    elif average_longitudinal_thickness_criteria == 'passed':

        # Average Measured Thickness from Critical Thickness Profiles based on the
        # circumferential CTP determined at the time of the inspection
        average_circumferential_thickness_criteria = g.check_average_circumferential_thickness_criteria(t_amC, FCA_ml,
                                                                                                        t_minL).result
        p = document.add_paragraph(g.check_average_circumferential_thickness_criteria(t_amC, FCA_ml,
                                                                                                        t_minL).text)
        p.style = math_font

        if average_circumferential_thickness_criteria == 'failed':

            print(average_circumferential_thickness_criteria)

        elif average_circumferential_thickness_criteria == 'passed':

            # step 10 ----------------------------------------------------------------
            p = document.add_paragraph('STEP 10')

            # Reduced MAWP of a conical or cylindrical shell based on the stresses in the circumferential or hoop direction, [bar]
            MAWP_rC = g.f_MAWP_rC(S, E, t_amS, FCA_ml, D_0, Y_B31).result
            p = document.add_paragraph(g.f_MAWP_rC(S, E, t_amS, FCA_ml, D_0, Y_B31).text)
            p.style = math_font

            # Reduced MAWP of a conical or cylindrical shell based on the stresses in the longitudinal direction, [bar]
            MAWP_rL = g.f_MAWP_rL(S, E, t_amC, FCA_ml, D_0, Y_B31).result
            p = document.add_paragraph(g.f_MAWP_rL(S, E, t_amC, FCA_ml, D_0, Y_B31).text)
            p.style = math_font

            #  MAWP criteria from Critical Thickness Profiles
            MAWP_criteria = g.check_MAWP_criteria(MAWP_rC, MAWP_rL, P).result
            p = document.add_paragraph(g.check_MAWP_criteria(MAWP_rC, MAWP_rL, P).text)
            p.style = math_font

            if MAWP_criteria == 'failed':

                print(MAWP_criteria)

            elif MAWP_criteria == 'passed':

                # step 11 ----------------------------------------------------------------
                p = document.add_paragraph('STEP 11')

                # Parameter which is needed for Minimum measured thickness criteria, [mm]
                t_lim = g.f_t_lim(t_nom).result
                p = document.add_paragraph(g.f_t_lim(t_nom).text)
                p.style = math_font

                # Minimum measured thickness criteria
                check_minimum_thickness_criteria = g.check_minimum_thickness_criteria(t_mm, FCA_ml, t_min, t_lim).result
                p = document.add_paragraph(g.check_minimum_thickness_criteria(t_mm, FCA_ml, t_min, t_lim).text)
                p.style = math_font

                if check_minimum_thickness_criteria == 'failed':

                    print(check_minimum_thickness_criteria)

                elif check_minimum_thickness_criteria == 'passed':

                    print(check_minimum_thickness_criteria)


# Save the document
document.save(os.path.join(directory, "test_report_GC.docx"))

if __name__ == '__main__':
    print('file is created')