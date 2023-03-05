from LocalCorrosionMethod import LocalCorrosionMethod as lc
import os
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

# Set the directory where the file should be saved
directory = r"..\reports"

# Create the document
document = Document()

# Set the font to one suitable for mathematical expressions
math_font = document.styles.add_style('MathFont', WD_STYLE_TYPE.PARAGRAPH)
math_font.font.name = 'Cambria Math'
math_font.font.size = Pt(12)

# input data
t = 5.49
pipe_type = 'Seamless'
M_ut = 12.5
LOSS = 0
FCA = 0
FCA_ml = 0
t_mm = 3.16
s = 50
S = 138
D_0 = 88.9
L_msd = 1000
g_r = 0
E = 1
MA = 0
Y_B31 = 0.4
t_sl = 0
RSF_a = 0.9
c = 50
EL = 1
EC = 1
P = 16.5
defect_type = 'Local metal loss'

input_data_text = """
Input data:
t = 5.49
pipe_type = 'Seamless'
M_ut = 12.5
LOSS = 0
FCA = 0
FCA_ml = 0
t_mm = 3.16
s = 50
S = 138
D_0 = 88.9
L_msd = 1000
g_r = 0
E = 1
MA = 0
Y_B31 = 0.4
t_sl = 0
RSF_a = 0.9
c = 50
EL = 1
EC = 1
P = 16.5
"""

paragraph = document.add_paragraph(input_data_text)
paragraph.style = math_font

# algorithm
# step 1.1 ------------------------------------------------
document.add_paragraph('STEP 1')
# Nominal or furnished thickness of the component adjusted for mill undertolerance as applicable, [mm]
t_nom = lc.f_t_nom(pipe_type, t, M_ut).result
document.add_paragraph(lc.f_t_nom(pipe_type, t, M_ut).text)

# Future corroded wall thickness away from the damage area, [mm]
t_c = lc.f_t_c(t_nom, LOSS, FCA).result
document.add_paragraph(lc.f_t_c(t_nom, LOSS, FCA).text)

# step 2 --------------------------------------------------
document.add_paragraph('STEP 2')
# Inside diameter of the cylinder, cone (at the location of the flaw), [mm]
D = lc.f_D(D_0, t_nom).result
document.add_paragraph(lc.f_D(D_0, t_nom).text)

# longitudinal flaw length parameter, [-]
λ = lc.f_λ(s, D, t_c).result
document.add_paragraph(lc.f_λ(s, D, t_c).text)

# step 2 --------------------------------------------------
document.add_paragraph('STEP 3')
# Remaining thickness ratio, [-]
R_t = lc.f_R_t(t_mm, FCA_ml, t_c).result
document.add_paragraph(lc.f_R_t(t_mm, FCA_ml, t_c).text)

if lc.check_R_t(R_t).result == 'failed':
    document.add_paragraph(lc.check_R_t(R_t).text)
else:
    document.add_paragraph(lc.check_R_t(R_t).text)

    if lc.check_thickness(t_mm, FCA_ml).result == 'failed':
        document.add_paragraph(lc.check_thickness(t_mm, FCA_ml).text)
    else:
        document.add_paragraph(lc.check_thickness(t_mm, FCA_ml).text)

        if lc.check_L_msd_criteria(L_msd, D, t_c).result == 'failed':
            document.add_paragraph(lc.check_L_msd_criteria(L_msd, D, t_c).text)
        else:
            document.add_paragraph(lc.check_L_msd_criteria(L_msd, D, t_c).text)

            # step 4 --------------------------------------------------
            document.add_paragraph('STEP 4')

            if defect_type == 'Groove' and lc.check_Groove_flaw_criteria(g_r, R_t, t_c).result == 'failed':
                document.add_paragraph(lc.check_Groove_flaw_criteria(g_r, R_t, t_c).text)
            elif (defect_type == 'Groove' and lc.check_Groove_flaw_criteria(g_r, R_t, t_c).result == 'passed') \
                                        or (defect_type == 'Local metal loss'):
                if defect_type == 'Groove' and lc.check_Groove_flaw_criteria(g_r, R_t, t_c).result == 'passed':
                    document.add_paragraph(lc.check_Groove_flaw_criteria(g_r, R_t, t_c).text)

                # Maximum allowable working pressure based on circumferential stress, [bar]
                MAWP_C = lc.f_MAWP_C(S, E, t_c, MA, D_0, Y_B31).result
                document.add_paragraph(lc.f_MAWP_C(S, E, t_c, MA, D_0, Y_B31).text)

                # Maximum allowable working pressure based on longitudinal stress, [bar]
                MAWP_L = lc.f_MAWP_L(S, E, t_c, MA, D_0, Y_B31, t_sl).result
                document.add_paragraph(lc.f_MAWP_L(S, E, t_c, MA, D_0, Y_B31, t_sl).text)

                MAWP = lc.f_MAWP(MAWP_C, MAWP_L).result
                document.add_paragraph(lc.f_MAWP(MAWP_C, MAWP_L).text)

                # step 6 --------------------------------------------------
                document.add_paragraph('STEP 6')

                # Folias factor based on the longitudinal extent of the LTA for a through-wall flaw, [-]
                M_t = lc.f_M_t(λ).result
                document.add_paragraph(lc.f_M_t(λ).text)

                # !!!!!-- if not acceptable then additional calculation ow MAWPr is needed ---!!!!
                # !!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!
                if lc.check_screening_criteria(λ, R_t, RSF_a, M_t).result == 'NOT acceptable':
                    document.add_paragraph(lc.check_screening_criteria(λ, R_t, RSF_a, M_t).text)

                    # Computed remaining strength factor based on the meridional extent of the LTA, [-]
                    RSF = lc.f_RSF(R_t, M_t).result
                    document.add_paragraph(lc.f_RSF(R_t, M_t).text)

                    # Reduced maximum allowable working pressure of the damaged component, [bar]
                    MAWPr = lc.f_MAWPr(MAWP, RSF, RSF_a).result
                    document.add_paragraph(lc.f_MAWPr(MAWP, RSF, RSF_a).text)

                else:
                    MAWPr = P
                    document.add_paragraph(f'As λ and R_t are acceptable MAWPr = P')

                # step 7 --------------------------------------------------
                document.add_paragraph('STEP 7')

                if lc.check_circumferential_extent_criteria(c, s, EL, EC) == 'failed':
                    document.add_paragraph(lc.check_circumferential_extent_criteria(c, s, EL, EC).text)
                else:
                    document.add_paragraph(lc.check_circumferential_extent_criteria(c, s, EL, EC).text)

                    #  Minimum required thickness for the component based on equipment design pressure or
                    #  equipment MAWP for longitudinal stresses
                    t_minL = lc.f_t_minL(MAWPr, D_0, S, E, P, Y_B31, t_sl, MA).result
                    document.add_paragraph(lc.f_t_minL(MAWPr, D_0, S, E, P, Y_B31, t_sl, MA).text)

                    if lc.check_minimum_thickness_required_criteria(t_minL, t_mm, FCA_ml).result == 'failed':
                        document.add_paragraph(lc.check_minimum_thickness_required_criteria(t_minL, t_mm, FCA_ml).text)

                        MAWPr_new = lc.f_MAWPr_new(MAWPr, t_mm, FCA_ml, t_minL).result
                        document.add_paragraph(lc.f_MAWPr_new(MAWPr, t_mm, FCA_ml, t_minL).text)
                    else:
                        document.add_paragraph(lc.check_minimum_thickness_required_criteria(t_minL, t_mm, FCA_ml).text)
                        print('all steps are passed')



# Save the document
document.save(os.path.join(directory, "test_report_LC.docx"))

if __name__ == '__main__':
    print('file is created')
